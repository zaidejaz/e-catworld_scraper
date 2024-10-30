import os
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt
from htmldocx import HtmlToDocx
from datetime import datetime
import re
import unicodedata
import json
from tqdm import tqdm

def extract_article_and_comments(html_file):
    try:
        with open(html_file, 'r', encoding='utf-8') as f:
            soup = BeautifulSoup(f.read(), 'html.parser')
        
        # Extract title
        title_elem = soup.select_one('h1.elementor-heading-title')
        title = title_elem.string.strip() if title_elem else "Untitled"
        
        # Extract date
        post_date = extract_date_from_element(soup)
        date_str = post_date.strftime('%B %d, %Y') if post_date else ""
        
        # Extract content
        content_elem = soup.select_one('.elementor-widget-theme-post-content')
        content = str(content_elem) if content_elem else ""
        
        # Extract comments
        comments = extract_comments(soup)
        
        return title, date_str, content, comments
    except Exception as e:
        print(f"Error in extract_article_and_comments: {str(e)}")
        raise

def extract_date_from_element(soup):
    date_selectors = [
        '.elementor-post-info__item--type-date',
        '.elementor-post-date',
        '.elementor-post-info time'
    ]
    
    for selector in date_selectors:
        date_elem = soup.select_one(selector)
        if date_elem:
            try:
                date_text = date_elem.get_text(strip=True)
                return datetime.strptime(date_text, '%B %d, %Y')
            except ValueError:
                continue
    return None

def extract_comments(soup):
    comments = []
    processed_ids = set()  # Track processed comment IDs
    comment_threads = soup.select('#disqus_thread li.post')
    
    for thread in comment_threads:
        try:
            # Get comment ID
            comment_id = thread.get('id', '')
            
            # Skip if already processed
            if comment_id in processed_ids:
                continue
                
            # Process main comment
            comment = extract_comment_data(thread)
            if comment:
                processed_ids.add(comment_id)
                
                # Process replies
                replies = []
                for reply in thread.select('.children li.post'):
                    reply_id = reply.get('id', '')
                    if reply_id not in processed_ids:
                        reply_data = extract_comment_data(reply)
                        if reply_data:
                            replies.append(reply_data)
                            processed_ids.add(reply_id)
                
                comment['replies'] = replies
                comments.append(comment)
                
        except Exception as e:
            print(f"Error processing comment: {str(e)}")
            continue
    
    return comments


def extract_comment_data(comment_elem):
    try:
        username_elem = comment_elem.select_one('.author')
        username = clean_username(username_elem.text) if username_elem else "Anonymous"
        
        date = comment_elem.select_one('a.time-ago')
        if not date:
            return None
        date = date['title']
        
        content = comment_elem.select_one('.post-message')
        if not content:
            return None
        content = process_html_content(str(content))
        
        if not content.strip():
            return None
            
        return {
            'username': username,
            'date': date,
            'content': content
        }
    except Exception as e:
        print(f"Error extracting comment data: {str(e)}")
        return None


def clean_username(username):
    cleaned = re.sub(r'\s+', ' ', username).strip()
    parts = cleaned.split()
    return ' '.join(sorted(set(parts), key=parts.index))

def process_html_content(html_content):
    soup = BeautifulSoup(html_content, 'html.parser')
    
    # Handle YouTube iframes and links
    for iframe in soup.find_all('iframe'):
        src = iframe.get('src', '')
        if 'youtube.com' in src or 'youtu.be' in src:
            iframe.replace_with(src)
        else:
            iframe.decompose()
    
    # Handle object tags with YouTube content
    for obj in soup.find_all('object'):
        params = obj.find_all('param', {'name': 'movie'})
        for param in params:
            value = param.get('value', '')
            if 'youtube.com' in value or 'youtu.be' in value:
                obj.replace_with(value)
                break
        obj.decompose()
    
    # Handle links
    for a in soup.find_all('a'):
        href = a.get('href', '')
        text = a.get_text(strip=True)
        
        if ('youtube.com' in href or 'youtu.be' in href):
            a.replace_with(href)
        elif 'e-catworld.com' in href:
            if text != href:
                a.replace_with(text)
            else:
                a.replace_with(href)
        else:
            if text:
                a.replace_with(text)
            else:
                a.decompose()
    
    # Remove all HTML tags but keep their text content
    text = soup.get_text(separator=' ', strip=True)
    
    # Clean up the text
    # Remove control characters
    text = ''.join(ch for ch in text if unicodedata.category(ch)[0] != "C")
    
    # Remove duplicate whitespace
    text = re.sub(r'\s+', ' ', text)
    
    # Remove references to disqus links
    text = re.sub(r'https?://disq\.us/\S+', '', text)
    
    # Clean up any empty parentheses and extra spaces
    text = re.sub(r'\(\s*\)', '', text)
    text = re.sub(r'\s+', ' ', text)
    
    # Handle paragraph breaks
    text = text.replace('[p]', '\n\n').replace('[/p]', '')
    
    # Final cleanup
    text = re.sub(r'\n\s*\n', '\n\n', text)
    return text.strip()

def process_html_files_by_year(input_dir, output_dir):
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    # Get all HTML files
    all_files = [f for f in os.listdir(input_dir) if f.endswith('.html')]
    articles_by_year = {}
    
    for filename in tqdm(all_files, desc="Processing files"):
        file_path = os.path.join(input_dir, filename)
        try:
            article_data = extract_article_and_comments(file_path)
            if article_data:
                title, date_str, content, comments = article_data
                try:
                    year = datetime.strptime(date_str, '%B %d, %Y').year if date_str else 0
                    if year not in articles_by_year:
                        articles_by_year[year] = []
                    articles_by_year[year].append(article_data)
                except ValueError:
                    print(f"Could not parse date for {filename}")
                    continue
                
        except Exception as e:
            print(f"Error processing {filename}: {str(e)}")
    
    # Create documents by year
    for year in sorted(articles_by_year.keys()):
        if year == 0:
            continue
        
        output_file = os.path.join(output_dir, f'e-catworld_articles_{year}.docx')
        create_word_document(output_file, articles_by_year[year])

def create_word_document(output_file, articles):
    doc = Document()
    
    # Sort articles by date
    articles.sort(key=lambda x: datetime.strptime(x[1], '%B %d, %Y') if x[1] else datetime.max)
    
    for article_data in articles:
        title, post_date, content, comments = article_data
        
        # Add title (font size 22)
        title_paragraph = doc.add_heading(level=1)
        title_run = title_paragraph.add_run(title)
        title_run.font.size = Pt(22)
        
        # Add post date
        if post_date:
            date_paragraph = doc.add_paragraph()
            date_run = date_paragraph.add_run(post_date)
            date_run.font.size = Pt(11)
        
        # Add content
        content_paragraph = doc.add_paragraph()
        content_run = content_paragraph.add_run(process_html_content(content))
        content_run.font.size = Pt(11)
        
        # Add comments
        if comments:
            doc.add_paragraph('---')
            for comment in comments:
                # Add comment
                comment_para = doc.add_paragraph()
                comment_para.add_run(f"{comment['username']} - {comment['date']}\n").bold = True
                comment_para.add_run(process_html_content(comment['content']))
                
                # Add replies
                if comment.get('replies'):
                    for reply in comment['replies']:
                        reply_para = doc.add_paragraph()
                        reply_para.add_run(f"{reply['username']} - {reply['date']}\n").italic = True
                        reply_para.add_run(process_html_content(reply['content']))
                
                doc.add_paragraph()
        
        doc.add_page_break()
    
    doc.save(output_file)

def add_comments_to_document(doc, comments):
    for comment in comments:
        # Add main comment
        comment_para = doc.add_paragraph()
        comment_para.add_run(f"{comment['username']} - {comment['date']}\n").bold = True
        comment_para.add_run(process_html_content(comment['content']))
        
        # Add replies if any
        if comment.get('replies'):
            for reply in comment['replies']:
                reply_para = doc.add_paragraph()
                reply_para.style = 'Quote'  # Use Quote style to indent replies
                reply_para.add_run(f"{reply['username']} - {reply['date']}\n").italic = True
                reply_para.add_run(process_html_content(reply['content']))
        
        # Add spacing between comment threads
        doc.add_paragraph()

if __name__ == "__main__":
    input_dir = "saved_html_pages"
    output_dir = "parsed_documents"
    process_html_files_by_year(input_dir, output_dir)